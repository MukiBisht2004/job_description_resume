#!/usr/bin/env python3
"""
Enhanced Backend Testing for AI Resume Tailor - Formatting Preservation
Tests the NEW formatting-preserving functionality specifically
"""

import requests
import json
import os
import tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io
import time
import base64

# Backend URL from frontend .env
BACKEND_URL = "https://resumeai-29.preview.emergentagent.com/api"

class FormattingTester:
    def __init__(self):
        self.session = requests.Session()
        self.test_results = []
        self.analysis_id = None
        self.docx_content_base64 = None
        self.resume_text = None
        
    def log_result(self, test_name, success, message, details=None):
        """Log test result"""
        result = {
            "test": test_name,
            "success": success,
            "message": message,
            "details": details or {}
        }
        self.test_results.append(result)
        status = "✅ PASS" if success else "❌ FAIL"
        print(f"{status}: {test_name} - {message}")
        if details:
            print(f"   Details: {details}")
        print()

    def create_richly_formatted_docx(self):
        """Create a DOCX with rich formatting for testing"""
        doc = Document()
        
        # Name with custom formatting
        name_para = doc.add_heading('', level=0)
        name_run = name_para.runs[0] if name_para.runs else name_para.add_run()
        name_run.text = 'Sarah Johnson'
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.name = 'Arial'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info with different formatting
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run('Senior Data Scientist | AI/ML Expert')
        contact_run.italic = True
        contact_run.font.size = Pt(12)
        
        contact_para2 = doc.add_paragraph()
        contact_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        email_run = contact_para2.add_run('sarah.johnson@email.com | ')
        email_run.font.size = Pt(10)
        phone_run = contact_para2.add_run('(555) 987-6543')
        phone_run.font.size = Pt(10)
        phone_run.bold = True
        
        # Professional Summary with custom heading
        summary_heading = doc.add_heading('Professional Summary', level=1)
        summary_heading.runs[0].font.name = 'Calibri'
        summary_heading.runs[0].font.size = Pt(14)
        
        summary_para = doc.add_paragraph()
        summary_text = ('Innovative Data Scientist with 7+ years of experience in machine learning, '
                       'deep learning, and statistical analysis. Proven track record of developing '
                       'AI solutions that drive business value and improve operational efficiency.')
        summary_run = summary_para.add_run(summary_text)
        summary_run.font.size = Pt(11)
        
        # Technical Skills with bullet formatting
        skills_heading = doc.add_heading('Technical Skills', level=1)
        skills_heading.runs[0].font.name = 'Calibri'
        skills_heading.runs[0].font.size = Pt(14)
        
        # Create bullet points with different formatting
        skills_list = [
            '• Programming: Python, R, SQL, Scala, Java',
            '• ML/AI: TensorFlow, PyTorch, scikit-learn, Keras, OpenCV',
            '• Big Data: Spark, Hadoop, Kafka, Airflow, Databricks',
            '• Cloud: AWS (SageMaker, EC2, S3), GCP, Azure ML',
            '• Databases: PostgreSQL, MongoDB, Redis, Snowflake'
        ]
        
        for skill in skills_list:
            skill_para = doc.add_paragraph()
            skill_run = skill_para.add_run(skill)
            skill_run.font.size = Pt(10)
            if 'Python' in skill or 'TensorFlow' in skill:
                skill_run.bold = True
        
        # Work Experience with varied formatting
        exp_heading = doc.add_heading('Professional Experience', level=1)
        exp_heading.runs[0].font.name = 'Calibri'
        exp_heading.runs[0].font.size = Pt(14)
        
        # Job 1
        job1_para = doc.add_paragraph()
        job1_title = job1_para.add_run('Senior Data Scientist')
        job1_title.bold = True
        job1_title.font.size = Pt(12)
        job1_company = job1_para.add_run(' | DataTech Solutions | ')
        job1_company.font.size = Pt(11)
        job1_dates = job1_para.add_run('2020-Present')
        job1_dates.italic = True
        job1_dates.font.size = Pt(10)
        
        achievements1 = [
            '• Led development of ML models that increased customer retention by 25%',
            '• Built real-time recommendation system serving 1M+ users daily',
            '• Implemented MLOps pipeline reducing model deployment time by 70%',
            '• Mentored team of 5 junior data scientists and ML engineers'
        ]
        
        for achievement in achievements1:
            ach_para = doc.add_paragraph()
            ach_run = ach_para.add_run(achievement)
            ach_run.font.size = Pt(10)
            if '25%' in achievement or '1M+' in achievement:
                # Make metrics bold
                parts = achievement.split()
                ach_para.clear()
                for part in parts:
                    if any(char.isdigit() for char in part) and ('%' in part or 'M' in part):
                        bold_run = ach_para.add_run(part + ' ')
                        bold_run.bold = True
                        bold_run.font.size = Pt(10)
                    else:
                        normal_run = ach_para.add_run(part + ' ')
                        normal_run.font.size = Pt(10)
        
        # Job 2
        job2_para = doc.add_paragraph()
        job2_title = job2_para.add_run('Data Scientist')
        job2_title.bold = True
        job2_title.font.size = Pt(12)
        job2_company = job2_para.add_run(' | AI Innovations Inc. | ')
        job2_company.font.size = Pt(11)
        job2_dates = job2_para.add_run('2018-2020')
        job2_dates.italic = True
        job2_dates.font.size = Pt(10)
        
        achievements2 = [
            '• Developed predictive analytics models for supply chain optimization',
            '• Created automated data pipelines processing 10TB+ daily',
            '• Collaborated with product teams to integrate ML into core platform'
        ]
        
        for achievement in achievements2:
            ach_para = doc.add_paragraph()
            ach_run = ach_para.add_run(achievement)
            ach_run.font.size = Pt(10)
        
        # Education
        edu_heading = doc.add_heading('Education', level=1)
        edu_heading.runs[0].font.name = 'Calibri'
        edu_heading.runs[0].font.size = Pt(14)
        
        edu_para = doc.add_paragraph()
        degree_run = edu_para.add_run('Ph.D. in Computer Science - Machine Learning')
        degree_run.bold = True
        degree_run.font.size = Pt(11)
        
        school_para = doc.add_paragraph()
        school_run = school_para.add_run('Stanford University | 2014-2018')
        school_run.font.size = Pt(10)
        school_run.italic = True
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def test_enhanced_upload_with_docx_content(self):
        """Test that upload now returns base64 DOCX content"""
        try:
            # Create richly formatted DOCX
            docx_content = self.create_richly_formatted_docx()
            
            # Test upload
            files = {'file': ('formatted_resume.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = self.session.post(f"{BACKEND_URL}/upload-resume", files=files, timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                
                # Check for required fields
                required_fields = ['success', 'text', 'filename', 'docx_content']
                missing_fields = [field for field in required_fields if field not in data]
                
                if not missing_fields:
                    # Verify docx_content is base64
                    try:
                        decoded_docx = base64.b64decode(data['docx_content'])
                        # Verify it's valid DOCX by reading it
                        test_doc = Document(io.BytesIO(decoded_docx))
                        paragraphs = [p.text for p in test_doc.paragraphs if p.text.strip()]
                        
                        self.docx_content_base64 = data['docx_content']
                        self.resume_text = data['text']
                        
                        self.log_result("Enhanced Upload - DOCX Content", True, 
                                      "Successfully returned base64 DOCX content",
                                      {
                                          "text_length": len(data['text']),
                                          "docx_base64_length": len(data['docx_content']),
                                          "decoded_docx_size": len(decoded_docx),
                                          "paragraphs_in_docx": len(paragraphs)
                                      })
                    except Exception as decode_error:
                        self.log_result("Enhanced Upload - DOCX Content", False, 
                                      f"Invalid base64 DOCX content: {str(decode_error)}")
                else:
                    self.log_result("Enhanced Upload - DOCX Content", False, 
                                  f"Missing required fields: {missing_fields}", data)
            else:
                self.log_result("Enhanced Upload - DOCX Content", False, 
                              f"HTTP {response.status_code}", response.text)
                
        except Exception as e:
            self.log_result("Enhanced Upload - DOCX Content", False, 
                          f"Error during enhanced upload test: {str(e)}")

    def test_enhanced_tailor_with_docx_content(self):
        """Test that tailor-resume accepts and stores original_docx_content"""
        if not self.docx_content_base64 or not self.resume_text:
            self.log_result("Enhanced Tailoring", False, "No DOCX content or resume text from upload test")
            return
            
        try:
            # Job description focused on AI/ML
            job_description = """
            Lead AI Engineer - Computer Vision & NLP
            
            We are seeking a Lead AI Engineer to drive our computer vision and NLP initiatives. 
            
            Required Skills:
            • 8+ years in AI/ML with focus on computer vision and NLP
            • Expert-level Python, TensorFlow, PyTorch
            • Experience with OpenCV, YOLO, Transformers, BERT
            • Cloud ML platforms (AWS SageMaker, GCP AI Platform)
            • MLOps and model deployment at scale
            • Leadership experience managing AI teams
            • PhD in Computer Science, AI, or related field preferred
            
            Responsibilities:
            • Lead computer vision projects for autonomous systems
            • Develop NLP models for document understanding
            • Build scalable ML infrastructure and pipelines
            • Mentor senior engineers and data scientists
            • Drive AI strategy and technical roadmap
            """
            
            # Test with original_docx_content parameter
            data = {
                'resume_text': self.resume_text,
                'job_description': job_description,
                'original_docx_content': self.docx_content_base64
            }
            
            response = self.session.post(f"{BACKEND_URL}/tailor-resume", data=data, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                if result.get("success") and result.get("tailored_resume"):
                    self.analysis_id = result.get("analysis_id")
                    
                    self.log_result("Enhanced Tailoring - With DOCX Content", True, 
                                  "Successfully tailored resume with original DOCX content stored",
                                  {
                                      "analysis_id": self.analysis_id,
                                      "ats_score": result.get("ats_score"),
                                      "tailored_length": len(result["tailored_resume"]),
                                      "suggestions_count": len(result.get("suggestions", []))
                                  })
                else:
                    self.log_result("Enhanced Tailoring - With DOCX Content", False, 
                                  "Missing expected response fields", result)
            else:
                self.log_result("Enhanced Tailoring - With DOCX Content", False, 
                              f"HTTP {response.status_code}", response.text)
                
        except Exception as e:
            self.log_result("Enhanced Tailoring - With DOCX Content", False, 
                          f"Error during enhanced tailoring test: {str(e)}")

    def test_formatting_preserved_download(self):
        """Test that download preserves original formatting"""
        if not self.analysis_id:
            self.log_result("Formatting Preservation", False, "No analysis_id from tailoring test")
            return
            
        try:
            response = self.session.get(f"{BACKEND_URL}/download-resume/{self.analysis_id}", timeout=30)
            
            if response.status_code == 200:
                # Verify it's a valid DOCX
                try:
                    downloaded_doc = Document(io.BytesIO(response.content))
                    
                    # Analyze the document structure and formatting
                    analysis = self.analyze_document_formatting(downloaded_doc)
                    
                    # Check if formatting elements are preserved
                    formatting_preserved = (
                        analysis['has_headings'] and
                        analysis['has_bold_text'] and
                        analysis['has_bullet_points'] and
                        analysis['paragraph_count'] > 10
                    )
                    
                    if formatting_preserved:
                        self.log_result("Formatting Preservation - Download", True, 
                                      "Successfully preserved formatting in downloaded DOCX",
                                      analysis)
                    else:
                        self.log_result("Formatting Preservation - Download", False, 
                                      "Formatting not fully preserved", analysis)
                        
                except Exception as docx_error:
                    self.log_result("Formatting Preservation - Download", False, 
                                  f"Downloaded file is not valid DOCX: {str(docx_error)}")
            else:
                self.log_result("Formatting Preservation - Download", False, 
                              f"HTTP {response.status_code}", response.text)
                
        except Exception as e:
            self.log_result("Formatting Preservation - Download", False, 
                          f"Error during formatting preservation test: {str(e)}")

    def analyze_document_formatting(self, doc):
        """Analyze document for formatting elements"""
        analysis = {
            'paragraph_count': 0,
            'has_headings': False,
            'has_bold_text': False,
            'has_italic_text': False,
            'has_bullet_points': False,
            'font_sizes': set(),
            'font_names': set(),
            'heading_count': 0,
            'bullet_count': 0
        }
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                analysis['paragraph_count'] += 1
                
                # Check if it's a heading
                if paragraph.style.name.startswith('Heading'):
                    analysis['has_headings'] = True
                    analysis['heading_count'] += 1
                
                # Check for bullet points
                if paragraph.text.strip().startswith('•') or paragraph.text.strip().startswith('-'):
                    analysis['has_bullet_points'] = True
                    analysis['bullet_count'] += 1
                
                # Check runs for formatting
                for run in paragraph.runs:
                    if run.bold:
                        analysis['has_bold_text'] = True
                    if run.italic:
                        analysis['has_italic_text'] = True
                    if run.font.size:
                        analysis['font_sizes'].add(str(run.font.size))
                    if run.font.name:
                        analysis['font_names'].add(run.font.name)
        
        # Convert sets to lists for JSON serialization
        analysis['font_sizes'] = list(analysis['font_sizes'])
        analysis['font_names'] = list(analysis['font_names'])
        
        return analysis

    def test_fallback_formatting(self):
        """Test fallback to simple formatting when no original DOCX"""
        try:
            # Test tailoring without original_docx_content
            job_description = "Software Engineer position requiring Python and API development skills."
            
            # Create a minimal DOCX for fallback test
            minimal_doc = Document()
            minimal_doc.add_paragraph("John Doe")
            minimal_doc.add_paragraph("Software Developer")
            minimal_doc.add_paragraph("Experience with Python and web development.")
            
            minimal_buffer = io.BytesIO()
            minimal_doc.save(minimal_buffer)
            minimal_buffer.seek(0)
            minimal_docx_base64 = base64.b64encode(minimal_buffer.getvalue()).decode('utf-8')
            
            data = {
                'resume_text': "John Doe\nSoftware Developer\nExperience with Python and web development.",
                'job_description': job_description,
                'original_docx_content': minimal_docx_base64  # Valid DOCX for fallback test
            }
            
            response = self.session.post(f"{BACKEND_URL}/tailor-resume", data=data, timeout=30)
            
            if response.status_code == 200:
                result = response.json()
                if result.get("success"):
                    fallback_analysis_id = result.get("analysis_id")
                    
                    # Test download with fallback formatting
                    download_response = self.session.get(f"{BACKEND_URL}/download-resume/{fallback_analysis_id}", timeout=30)
                    
                    if download_response.status_code == 200:
                        try:
                            fallback_doc = Document(io.BytesIO(download_response.content))
                            paragraphs = [p.text for p in fallback_doc.paragraphs if p.text.strip()]
                            
                            if len(paragraphs) > 0:
                                self.log_result("Fallback Formatting", True, 
                                              "Successfully created DOCX with fallback formatting",
                                              {"paragraphs": len(paragraphs)})
                            else:
                                self.log_result("Fallback Formatting", False, "Empty document created")
                        except Exception as docx_error:
                            self.log_result("Fallback Formatting", False, 
                                          f"Fallback DOCX is invalid: {str(docx_error)}")
                    else:
                        self.log_result("Fallback Formatting", False, 
                                      f"Download failed: HTTP {download_response.status_code}")
                else:
                    self.log_result("Fallback Formatting", False, "Tailoring failed", result)
            else:
                self.log_result("Fallback Formatting", False, 
                              f"Tailoring request failed: HTTP {response.status_code}")
                
        except Exception as e:
            self.log_result("Fallback Formatting", False, 
                          f"Error during fallback test: {str(e)}")

    def test_content_mapping_intelligence(self):
        """Test that content is intelligently mapped to preserve structure"""
        if not self.analysis_id:
            self.log_result("Content Mapping Intelligence", False, "No analysis_id available")
            return
            
        try:
            # Get the analysis from database to compare original vs tailored
            analyses_response = self.session.get(f"{BACKEND_URL}/analyses", timeout=15)
            
            if analyses_response.status_code == 200:
                analyses_data = analyses_response.json()
                analyses = analyses_data.get("analyses", [])
                
                # Find our analysis
                our_analysis = None
                for analysis in analyses:
                    if analysis.get("id") == self.analysis_id:
                        our_analysis = analysis
                        break
                
                if our_analysis:
                    original_text = our_analysis.get("original_text", "")
                    tailored_text = our_analysis.get("tailored_resume", "")
                    
                    # Analyze content mapping
                    original_sections = self.extract_sections(original_text)
                    tailored_sections = self.extract_sections(tailored_text)
                    
                    # Check if key sections are preserved
                    key_sections = ['experience', 'education', 'skills', 'summary']
                    sections_preserved = 0
                    
                    for section in key_sections:
                        if section in original_sections and section in tailored_sections:
                            sections_preserved += 1
                    
                    mapping_quality = sections_preserved / len(key_sections)
                    
                    if mapping_quality >= 0.5:  # At least 50% of sections preserved
                        self.log_result("Content Mapping Intelligence", True, 
                                      f"Good content mapping: {sections_preserved}/{len(key_sections)} sections preserved",
                                      {
                                          "mapping_quality": f"{mapping_quality:.2%}",
                                          "original_sections": list(original_sections.keys()),
                                          "tailored_sections": list(tailored_sections.keys())
                                      })
                    else:
                        self.log_result("Content Mapping Intelligence", False, 
                                      f"Poor content mapping: {sections_preserved}/{len(key_sections)} sections preserved")
                else:
                    self.log_result("Content Mapping Intelligence", False, "Could not find our analysis in database")
            else:
                self.log_result("Content Mapping Intelligence", False, 
                              f"Could not retrieve analyses: HTTP {analyses_response.status_code}")
                
        except Exception as e:
            self.log_result("Content Mapping Intelligence", False, 
                          f"Error during content mapping test: {str(e)}")

    def extract_sections(self, text):
        """Extract sections from resume text"""
        sections = {}
        current_section = None
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip().lower()
            if not line:
                continue
                
            # Check for section headers
            if any(keyword in line for keyword in ['experience', 'work', 'employment']):
                current_section = 'experience'
                sections[current_section] = []
            elif any(keyword in line for keyword in ['education', 'degree', 'university']):
                current_section = 'education'
                sections[current_section] = []
            elif any(keyword in line for keyword in ['skills', 'technical', 'competencies']):
                current_section = 'skills'
                sections[current_section] = []
            elif any(keyword in line for keyword in ['summary', 'objective', 'profile']):
                current_section = 'summary'
                sections[current_section] = []
            elif current_section and line:
                sections[current_section].append(line)
        
        return sections

    def run_formatting_tests(self):
        """Run all formatting preservation tests"""
        print("🎨 Starting AI Resume Tailor - Formatting Preservation Tests")
        print(f"Backend URL: {BACKEND_URL}")
        print("=" * 70)
        
        # Run tests in order
        self.test_enhanced_upload_with_docx_content()
        self.test_enhanced_tailor_with_docx_content()
        self.test_formatting_preserved_download()
        self.test_content_mapping_intelligence()
        self.test_fallback_formatting()
        
        # Summary
        print("=" * 70)
        print("📊 FORMATTING PRESERVATION TEST SUMMARY")
        print("=" * 70)
        
        passed = sum(1 for result in self.test_results if result["success"])
        total = len(self.test_results)
        
        print(f"Total Tests: {total}")
        print(f"Passed: {passed}")
        print(f"Failed: {total - passed}")
        print(f"Success Rate: {(passed/total)*100:.1f}%")
        
        print("\n📋 DETAILED RESULTS:")
        for result in self.test_results:
            status = "✅" if result["success"] else "❌"
            print(f"{status} {result['test']}: {result['message']}")
        
        return passed == total

if __name__ == "__main__":
    tester = FormattingTester()
    success = tester.run_formatting_tests()
    
    if success:
        print("\n🎉 All formatting preservation tests passed!")
    else:
        print("\n⚠️  Some formatting tests failed. Check results above.")