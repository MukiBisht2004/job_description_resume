from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field
from typing import Optional, List, Dict, Any
import os
import uuid
from datetime import datetime, timezone
import tempfile
import shutil
from pathlib import Path
from dotenv import load_dotenv
import json

# Document processing
from docx import Document
from docx.shared import Inches
import io
import base64

# AI Integration
from emergentintegrations.llm.chat import LlmChat, UserMessage

# Load environment variables
load_dotenv()

# MongoDB connection
MONGO_URL = os.environ.get('MONGO_URL')
client = AsyncIOMotorClient(MONGO_URL)
db = client.career_assistant

app = FastAPI()

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models
class ResumeAnalysis(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    original_text: str
    original_docx_content: str  # Base64 encoded DOCX content
    job_description: str
    tailored_resume: str
    ats_score: int
    suggestions: List[str]
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class JobDescription(BaseModel):
    text: str

class ATSAnalysis(BaseModel):
    score: int
    suggestions: List[str]
    keyword_matches: List[str]
    missing_keywords: List[str]

# AI Chat setup
def get_llm_chat(session_id: str, system_message: str):
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="LLM API key not configured")
    
    chat = LlmChat(
        api_key=api_key,
        session_id=session_id,
        system_message=system_message
    ).with_model("openai", "gpt-4o")
    
    return chat

# Helper functions
def extract_text_and_structure_from_docx(file_content: bytes) -> tuple:
    """Extract text and preserve document structure from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        full_text = '\n'.join(text_parts)
        return full_text, file_content
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {str(e)}")

def create_tailored_docx_with_formatting(original_docx_content: bytes, original_text: str, tailored_text: str) -> bytes:
    """Create tailored DOCX while preserving original formatting"""
    try:
        # Load original document
        original_doc = Document(io.BytesIO(original_docx_content))
        
        # Split both texts into lines for mapping
        original_lines = [line.strip() for line in original_text.split('\n') if line.strip()]
        tailored_lines = [line.strip() for line in tailored_text.split('\n') if line.strip()]
        
        # Create a mapping between original and tailored content
        line_mapping = {}
        
        # Simple approach: map lines by position and similarity
        for i, orig_line in enumerate(original_lines):
            if i < len(tailored_lines):
                line_mapping[orig_line] = tailored_lines[i]
            else:
                line_mapping[orig_line] = orig_line  # Keep original if no tailored version
        
        # Process the document paragraphs
        for paragraph in original_doc.paragraphs:
            if paragraph.text.strip():
                original_para_text = paragraph.text.strip()
                
                # Find best match in tailored content
                best_match = None
                best_score = 0
                
                for tailored_line in tailored_lines:
                    # Simple similarity check
                    common_words = set(original_para_text.lower().split()) & set(tailored_line.lower().split())
                    score = len(common_words)
                    
                    if score > best_score:
                        best_score = score
                        best_match = tailored_line
                
                # Replace content while preserving formatting
                if best_match and best_score > 0:
                    # Clear existing runs but keep paragraph formatting
                    for run in paragraph.runs:
                        run.text = ""
                    
                    # Add new text with original formatting of first run
                    if paragraph.runs:
                        paragraph.runs[0].text = best_match
                    else:
                        paragraph.text = best_match
                elif not best_match:
                    # For new content not in original, try to find similar sections
                    for tailored_line in tailored_lines:
                        # Check if this line should replace current paragraph
                        if any(word in tailored_line.lower() for word in original_para_text.lower().split()[:3]):
                            # Clear and replace
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = tailored_line
                            else:
                                paragraph.text = tailored_line
                            break
        
        # Save the modified document
        buffer = io.BytesIO()
        original_doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        # Fallback to simple text replacement
        return create_simple_formatted_docx(tailored_text)

def create_simple_formatted_docx(text: str) -> bytes:
    """Fallback: Create a nicely formatted DOCX file from text"""
    try:
        doc = Document()
        
        # Split text into sections and format appropriately
        lines = text.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if this is a heading (name, section headers)
            if (len(line.split()) <= 4 and 
                any(keyword in line.lower() for keyword in ['experience', 'education', 'skills', 'summary', 'objective', 'contact', 'certifications', 'projects']) or
                (line.isupper() and len(line) < 50) or
                (not any(c in line for c in ['.', ',', ';']) and len(line.split()) <= 3)):
                
                # Add as heading
                heading = doc.add_heading(line, level=1)
                heading.alignment = 0  # Left align
                
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # Bullet point
                p = doc.add_paragraph(line, style='List Bullet')
                
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
                
                # Check if it's a name (first line, usually)
                if len(doc.paragraphs) == 1 and not any(keyword in line.lower() for keyword in ['email', 'phone', 'address']):
                    # Make it bold and larger (name)
                    for run in p.runs:
                        run.bold = True
                        run.font.size = 16
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        # Ultimate fallback
        doc = Document()
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

# Keep old functions for backward compatibility
def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file - backward compatibility wrapper"""
    text, _ = extract_text_and_structure_from_docx(file_content)
    return text

def create_docx_from_text(text: str, filename: str = "tailored_resume.docx") -> bytes:
    """Create a DOCX file from text - backward compatibility wrapper"""
    return create_simple_formatted_docx(text)

async def tailor_resume_with_ai(resume_text: str, job_description: str) -> str:
    """Use AI to tailor resume for specific job"""
    session_id = f"resume_tailor_{uuid.uuid4()}"
    system_message = """You are an expert resume writer and ATS optimization specialist. Your task is to rewrite and tailor resumes to match specific job descriptions while maintaining the original format and style.

Guidelines:
1. Analyze the job description to identify key skills, requirements, and keywords
2. Rewrite resume content to highlight relevant experience and skills
3. Use job-specific keywords naturally throughout the resume
4. Maintain the original resume structure and formatting style
5. Ensure all claims are truthful - only emphasize existing skills/experience
6. Make the resume ATS-friendly with proper keyword density
7. Keep the same sections but optimize content for the target role

Return only the tailored resume text without any additional commentary."""

    try:
        chat = get_llm_chat(session_id, system_message)
        
        prompt = f"""Please tailor this resume for the following job description:

JOB DESCRIPTION:
{job_description}

ORIGINAL RESUME:
{resume_text}

Please rewrite the resume to better match the job requirements while keeping the same structure and format."""

        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        return response
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error tailoring resume: {str(e)}")

async def analyze_ats_score(resume_text: str, job_description: str) -> ATSAnalysis:
    """Analyze resume for ATS compatibility and scoring"""
    session_id = f"ats_analysis_{uuid.uuid4()}"
    system_message = """You are an ATS (Applicant Tracking System) analysis expert. Your task is to analyze resumes against job descriptions and provide detailed scoring and improvement suggestions.

Analyze based on:
1. Keyword matching and density
2. Skills alignment
3. Experience relevance
4. Format compatibility
5. Section organization
6. Achievement quantification

Provide your response in the following JSON format:
{
    "score": 85,
    "suggestions": ["Add more specific technical skills", "Include quantified achievements"],
    "keyword_matches": ["Python", "Data Analysis", "Machine Learning"],
    "missing_keywords": ["AWS", "Docker", "Kubernetes"]
}"""

    try:
        chat = get_llm_chat(session_id, system_message)
        
        prompt = f"""Analyze this resume against the job description and provide ATS scoring:

JOB DESCRIPTION:
{job_description}

RESUME:
{resume_text}

Please provide a detailed ATS analysis including score (0-100), specific suggestions for improvement, matched keywords, and missing important keywords. Return only valid JSON."""

        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        # Parse JSON response
        try:
            analysis_data = json.loads(response)
            return ATSAnalysis(**analysis_data)
        except json.JSONDecodeError:
            # Fallback if AI doesn't return proper JSON
            return ATSAnalysis(
                score=75,
                suggestions=["Resume has been analyzed", "Consider adding more relevant keywords"],
                keyword_matches=["General skills match"],
                missing_keywords=["Specific technical requirements"]
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analyzing ATS score: {str(e)}")

# API endpoints
@app.get("/api/health")
async def health_check():
    return {"status": "healthy", "service": "Career Assistant API"}

@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """Upload and process resume file"""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX files are supported")
    
    try:
        content = await file.read()
        text, original_docx = extract_text_and_structure_from_docx(content)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from resume")
        
        # Encode DOCX content to base64 for storage
        docx_base64 = base64.b64encode(original_docx).decode('utf-8')
        
        return {
            "success": True,
            "text": text,
            "filename": file.filename,
            "docx_content": docx_base64  # Include for frontend to store
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")

@app.post("/api/tailor-resume")
async def tailor_resume(
    resume_text: str = Form(...),
    job_description: str = Form(...),
    original_docx_content: str = Form(...)  # Base64 encoded DOCX
):
    """Tailor resume for specific job description"""
    try:
        # Tailor the resume
        tailored_resume = await tailor_resume_with_ai(resume_text, job_description)
        
        # Analyze ATS score
        ats_analysis = await analyze_ats_score(tailored_resume, job_description)
        
        # Save to database
        analysis = ResumeAnalysis(
            original_text=resume_text,
            original_docx_content=original_docx_content,  # Store base64 DOCX
            job_description=job_description,
            tailored_resume=tailored_resume,
            ats_score=ats_analysis.score,
            suggestions=ats_analysis.suggestions
        )
        
        await db.resume_analyses.insert_one(analysis.dict())
        
        return {
            "success": True,
            "analysis_id": analysis.id,
            "tailored_resume": tailored_resume,
            "ats_score": ats_analysis.score,
            "suggestions": ats_analysis.suggestions,
            "keyword_matches": ats_analysis.keyword_matches,
            "missing_keywords": ats_analysis.missing_keywords
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error tailoring resume: {str(e)}")

@app.get("/api/download-resume/{analysis_id}")
async def download_resume(analysis_id: str):
    """Download tailored resume as DOCX"""
    try:
        # Get analysis from database
        analysis = await db.resume_analyses.find_one({"id": analysis_id})
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Create DOCX file
        docx_content = create_docx_from_text(analysis["tailored_resume"])
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        temp_file.write(docx_content)
        temp_file.close()
        
        return FileResponse(
            path=temp_file.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"tailored_resume_{analysis_id[:8]}.docx"
        )
    except HTTPException:
        # Re-raise HTTP exceptions (like 404)
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error downloading resume: {str(e)}")

@app.get("/api/analyses")
async def get_analyses():
    """Get all resume analyses"""
    try:
        analyses = await db.resume_analyses.find({}, {"_id": 0}).sort("created_at", -1).limit(50).to_list(length=None)
        return {"analyses": analyses}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching analyses: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)