#!/usr/bin/env python3
"""
Comprehensive Backend API Testing for AI Resume Tailor
Tests all backend endpoints and functionality
"""

import requests
import json
import os
import tempfile
from docx import Document
import io
import time

# Backend URL from frontend .env
BACKEND_URL = "https://resumeai-29.preview.emergentagent.com/api"

class BackendTester:
    def __init__(self):
        self.session = requests.Session()
        self.test_results = []
        self.analysis_id = None
        
    def log_result(self, test_name, success, message, details=None):
        """Log test result"""
        result = {
            "test": test_name,
            "success": success,
            "message": message,
            "details": details or {}
        }
        self.test_results.append(result)
        status = "✅ PASS" if success else "❌ FAIL"
        print(f"{status}: {test_name} - {message}")
        if details:
            print(f"   Details: {details}")
        print()

    def create_sample_docx(self):
        """Create a sample DOCX resume for testing"""
        doc = Document()
        
        # Add resume content
        doc.add_heading('John Smith', 0)
        doc.add_paragraph('Software Engineer | Full Stack Developer')
        doc.add_paragraph('Email: john.smith@email.com | Phone: (555) 123-4567')
        
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph('Experienced software engineer with 5+ years in full-stack development. '
                         'Proficient in Python, JavaScript, React, and cloud technologies. '
                         'Strong background in API development, database design, and agile methodologies.')
        
        doc.add_heading('Technical Skills', level=1)
        doc.add_paragraph('• Programming Languages: Python, JavaScript, Java, SQL')
        doc.add_paragraph('• Frameworks: React, FastAPI, Django, Node.js')
        doc.add_paragraph('• Databases: MongoDB, PostgreSQL, MySQL')
        doc.add_paragraph('• Cloud: AWS, Docker, Kubernetes')
        doc.add_paragraph('• Tools: Git, Jenkins, JIRA, Postman')
        
        doc.add_heading('Work Experience', level=1)
        doc.add_paragraph('Senior Software Engineer | TechCorp Inc. | 2021-Present')
        doc.add_paragraph('• Developed and maintained REST APIs serving 100K+ daily requests')
        doc.add_paragraph('• Built responsive web applications using React and TypeScript')
        doc.add_paragraph('• Implemented CI/CD pipelines reducing deployment time by 60%')
        doc.add_paragraph('• Collaborated with cross-functional teams in agile environment')
        
        doc.add_paragraph('Software Developer | StartupXYZ | 2019-2021')
        doc.add_paragraph('• Created full-stack web applications using Python and JavaScript')
        doc.add_paragraph('• Designed and optimized database schemas for improved performance')
        doc.add_paragraph('• Participated in code reviews and mentored junior developers')
        
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Bachelor of Science in Computer Science')
        doc.add_paragraph('University of Technology | 2015-2019')
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def test_health_endpoint(self):
        """Test GET /api/health endpoint"""
        try:
            response = self.session.get(f"{BACKEND_URL}/health", timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                if data.get("status") == "healthy":
                    self.log_result("Health Check", True, "API is healthy and responding")
                else:
                    self.log_result("Health Check", False, "Unexpected health response", data)
            else:
                self.log_result("Health Check", False, f"HTTP {response.status_code}", response.text)
                
        except Exception as e:
            self.log_result("Health Check", False, f"Connection error: {str(e)}")

    def test_upload_resume(self):
        """Test POST /api/upload-resume endpoint"""
        try:
            # Create sample DOCX
            docx_content = self.create_sample_docx()
            
            # Test valid DOCX upload
            files = {'file': ('test_resume.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = self.session.post(f"{BACKEND_URL}/upload-resume", files=files, timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                if data.get("success") and data.get("text"):
                    self.log_result("Resume Upload - Valid DOCX", True, "Successfully uploaded and extracted text", 
                                  {"text_length": len(data["text"]), "filename": data.get("filename")})
                    # Store extracted text for later tests
                    self.resume_text = data["text"]
                else:
                    self.log_result("Resume Upload - Valid DOCX", False, "Missing expected response fields", data)
            else:
                self.log_result("Resume Upload - Valid DOCX", False, f"HTTP {response.status_code}", response.text)
            
            # Test invalid file type
            files = {'file': ('test.txt', b'This is not a DOCX file', 'text/plain')}
            response = self.session.post(f"{BACKEND_URL}/upload-resume", files=files, timeout=10)
            
            if response.status_code == 400:
                self.log_result("Resume Upload - Invalid File Type", True, "Correctly rejected non-DOCX file")
            else:
                self.log_result("Resume Upload - Invalid File Type", False, f"Should reject non-DOCX files, got HTTP {response.status_code}")
                
        except Exception as e:
            self.log_result("Resume Upload", False, f"Error during upload test: {str(e)}")

    def test_tailor_resume(self):
        """Test POST /api/tailor-resume endpoint"""
        if not hasattr(self, 'resume_text'):
            self.log_result("Resume Tailoring", False, "No resume text available from upload test")
            return
            
        try:
            # Sample job description
            job_description = """
            Senior Python Developer - AI/ML Focus
            
            We are seeking an experienced Python developer to join our AI/ML team. The ideal candidate will have:
            
            Required Skills:
            • 5+ years of Python development experience
            • Strong experience with FastAPI, Django, or Flask
            • Machine Learning frameworks (TensorFlow, PyTorch, scikit-learn)
            • Database experience with MongoDB and PostgreSQL
            • Cloud platforms (AWS, GCP, Azure)
            • Docker and Kubernetes experience
            • REST API development and integration
            • Git version control and CI/CD pipelines
            
            Responsibilities:
            • Develop and maintain ML pipelines and APIs
            • Build scalable backend services for AI applications
            • Collaborate with data scientists and ML engineers
            • Optimize model performance and deployment
            • Implement automated testing and monitoring
            
            Nice to have:
            • Experience with LLM integration (OpenAI, Anthropic)
            • Knowledge of vector databases
            • MLOps experience
            • Agile/Scrum methodology experience
            """
            
            # Test resume tailoring
            data = {
                'resume_text': self.resume_text,
                'job_description': job_description
            }
            
            response = self.session.post(f"{BACKEND_URL}/tailor-resume", data=data, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                if result.get("success") and result.get("tailored_resume"):
                    self.analysis_id = result.get("analysis_id")
                    ats_score = result.get("ats_score", 0)
                    suggestions = result.get("suggestions", [])
                    keyword_matches = result.get("keyword_matches", [])
                    missing_keywords = result.get("missing_keywords", [])
                    
                    self.log_result("Resume Tailoring - AI Processing", True, 
                                  f"Successfully tailored resume with ATS score: {ats_score}",
                                  {
                                      "analysis_id": self.analysis_id,
                                      "ats_score": ats_score,
                                      "suggestions_count": len(suggestions),
                                      "keyword_matches_count": len(keyword_matches),
                                      "missing_keywords_count": len(missing_keywords),
                                      "tailored_resume_length": len(result["tailored_resume"])
                                  })
                    
                    # Validate ATS score is reasonable
                    if 0 <= ats_score <= 100:
                        self.log_result("ATS Scoring", True, f"Valid ATS score: {ats_score}")
                    else:
                        self.log_result("ATS Scoring", False, f"Invalid ATS score: {ats_score}")
                        
                    # Validate suggestions exist
                    if suggestions and len(suggestions) > 0:
                        self.log_result("ATS Suggestions", True, f"Generated {len(suggestions)} suggestions")
                    else:
                        self.log_result("ATS Suggestions", False, "No suggestions generated")
                        
                else:
                    self.log_result("Resume Tailoring", False, "Missing expected response fields", result)
            else:
                self.log_result("Resume Tailoring", False, f"HTTP {response.status_code}", response.text)
                
        except Exception as e:
            self.log_result("Resume Tailoring", False, f"Error during tailoring test: {str(e)}")

    def test_download_resume(self):
        """Test GET /api/download-resume/{analysis_id} endpoint"""
        if not self.analysis_id:
            self.log_result("Resume Download", False, "No analysis_id available from tailoring test")
            return
            
        try:
            response = self.session.get(f"{BACKEND_URL}/download-resume/{self.analysis_id}", timeout=30)
            
            if response.status_code == 200:
                # Check if response is DOCX file
                content_type = response.headers.get('content-type', '')
                if 'wordprocessingml' in content_type or 'docx' in content_type:
                    # Verify it's actually a valid DOCX by trying to read it
                    try:
                        doc = Document(io.BytesIO(response.content))
                        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                        self.log_result("Resume Download", True, 
                                      f"Successfully downloaded DOCX file ({len(response.content)} bytes)",
                                      {"paragraphs_count": len(paragraphs), "content_type": content_type})
                    except Exception as docx_error:
                        self.log_result("Resume Download", False, f"Downloaded file is not valid DOCX: {str(docx_error)}")
                else:
                    self.log_result("Resume Download", False, f"Unexpected content type: {content_type}")
            else:
                self.log_result("Resume Download", False, f"HTTP {response.status_code}", response.text)
                
            # Test invalid analysis_id
            response = self.session.get(f"{BACKEND_URL}/download-resume/invalid-id", timeout=10)
            if response.status_code == 404:
                self.log_result("Resume Download - Invalid ID", True, "Correctly returned 404 for invalid analysis_id")
            else:
                self.log_result("Resume Download - Invalid ID", False, f"Should return 404 for invalid ID, got {response.status_code}")
                
        except Exception as e:
            self.log_result("Resume Download", False, f"Error during download test: {str(e)}")

    def test_get_analyses(self):
        """Test GET /api/analyses endpoint"""
        try:
            response = self.session.get(f"{BACKEND_URL}/analyses", timeout=15)
            
            if response.status_code == 200:
                data = response.json()
                analyses = data.get("analyses", [])
                
                if isinstance(analyses, list):
                    # Check if our analysis is in the list
                    found_our_analysis = False
                    if self.analysis_id:
                        for analysis in analyses:
                            if analysis.get("id") == self.analysis_id:
                                found_our_analysis = True
                                break
                    
                    if found_our_analysis or len(analyses) > 0:
                        self.log_result("MongoDB Storage - Get Analyses", True, 
                                      f"Successfully retrieved {len(analyses)} analyses",
                                      {"found_test_analysis": found_our_analysis})
                    else:
                        self.log_result("MongoDB Storage - Get Analyses", True, 
                                      "Endpoint works but no analyses found (expected for fresh database)")
                else:
                    self.log_result("MongoDB Storage - Get Analyses", False, "Response is not a list", data)
            else:
                self.log_result("MongoDB Storage - Get Analyses", False, f"HTTP {response.status_code}", response.text)
                
        except Exception as e:
            self.log_result("MongoDB Storage - Get Analyses", False, f"Error during analyses test: {str(e)}")

    def test_error_handling(self):
        """Test various error conditions"""
        try:
            # Test missing resume text in tailor-resume
            data = {'job_description': 'Test job description'}
            response = self.session.post(f"{BACKEND_URL}/tailor-resume", data=data, timeout=10)
            
            if response.status_code == 422:  # FastAPI validation error
                self.log_result("Error Handling - Missing Resume Text", True, "Correctly validated missing resume_text")
            else:
                self.log_result("Error Handling - Missing Resume Text", False, f"Expected 422, got {response.status_code}")
            
            # Test missing job description in tailor-resume
            data = {'resume_text': 'Test resume text'}
            response = self.session.post(f"{BACKEND_URL}/tailor-resume", data=data, timeout=10)
            
            if response.status_code == 422:  # FastAPI validation error
                self.log_result("Error Handling - Missing Job Description", True, "Correctly validated missing job_description")
            else:
                self.log_result("Error Handling - Missing Job Description", False, f"Expected 422, got {response.status_code}")
                
        except Exception as e:
            self.log_result("Error Handling", False, f"Error during error handling test: {str(e)}")

    def run_all_tests(self):
        """Run all backend tests"""
        print("🚀 Starting AI Resume Tailor Backend API Tests")
        print(f"Backend URL: {BACKEND_URL}")
        print("=" * 60)
        
        # Run tests in order
        self.test_health_endpoint()
        self.test_upload_resume()
        self.test_tailor_resume()
        self.test_download_resume()
        self.test_get_analyses()
        self.test_error_handling()
        
        # Summary
        print("=" * 60)
        print("📊 TEST SUMMARY")
        print("=" * 60)
        
        passed = sum(1 for result in self.test_results if result["success"])
        total = len(self.test_results)
        
        print(f"Total Tests: {total}")
        print(f"Passed: {passed}")
        print(f"Failed: {total - passed}")
        print(f"Success Rate: {(passed/total)*100:.1f}%")
        
        print("\n📋 DETAILED RESULTS:")
        for result in self.test_results:
            status = "✅" if result["success"] else "❌"
            print(f"{status} {result['test']}: {result['message']}")
        
        return passed == total

if __name__ == "__main__":
    tester = BackendTester()
    success = tester.run_all_tests()
    
    if success:
        print("\n🎉 All tests passed! Backend API is working correctly.")
    else:
        print("\n⚠️  Some tests failed. Check the results above for details.")