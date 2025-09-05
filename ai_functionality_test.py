#!/usr/bin/env python3
"""
Additional AI Functionality Tests for Resume Tailor
Tests the AI tailoring quality and different job scenarios
"""

import requests
import json
from docx import Document
import io

BACKEND_URL = "https://resumeai-29.preview.emergentagent.com/api"

def create_marketing_resume():
    """Create a marketing-focused resume"""
    doc = Document()
    
    doc.add_heading('Sarah Johnson', 0)
    doc.add_paragraph('Digital Marketing Manager | Brand Strategist')
    doc.add_paragraph('Email: sarah.johnson@email.com | Phone: (555) 987-6543')
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('Creative marketing professional with 6+ years of experience in digital marketing, '
                     'brand management, and campaign optimization. Proven track record of increasing '
                     'brand awareness by 150% and driving 40% growth in lead generation.')
    
    doc.add_heading('Core Competencies', level=1)
    doc.add_paragraph('• Digital Marketing Strategy & Campaign Management')
    doc.add_paragraph('• Social Media Marketing (Facebook, Instagram, LinkedIn, Twitter)')
    doc.add_paragraph('• Content Creation & Copywriting')
    doc.add_paragraph('• SEO/SEM & Google Analytics')
    doc.add_paragraph('• Email Marketing & Marketing Automation')
    doc.add_paragraph('• Brand Development & Market Research')
    
    doc.add_heading('Professional Experience', level=1)
    doc.add_paragraph('Digital Marketing Manager | BrandCorp | 2020-Present')
    doc.add_paragraph('• Developed and executed integrated marketing campaigns resulting in 40% increase in qualified leads')
    doc.add_paragraph('• Managed social media presence across 5 platforms with 200K+ combined followers')
    doc.add_paragraph('• Implemented marketing automation workflows improving conversion rates by 25%')
    doc.add_paragraph('• Collaborated with design team to create compelling visual content and brand materials')
    
    doc.add_paragraph('Marketing Specialist | StartupMedia | 2018-2020')
    doc.add_paragraph('• Created content marketing strategy that increased website traffic by 300%')
    doc.add_paragraph('• Managed Google Ads campaigns with $50K monthly budget and 4.2% CTR')
    doc.add_paragraph('• Conducted market research and competitor analysis for product launches')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def test_ai_tailoring_scenarios():
    """Test AI tailoring with different job scenarios"""
    session = requests.Session()
    
    # Test 1: Marketing to Tech transition
    print("🧪 Testing AI Tailoring: Marketing to Tech Transition")
    print("=" * 60)
    
    # Upload marketing resume
    marketing_resume = create_marketing_resume()
    files = {'file': ('marketing_resume.docx', marketing_resume, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    response = session.post(f"{BACKEND_URL}/upload-resume", files=files, timeout=30)
    
    if response.status_code != 200:
        print("❌ Failed to upload marketing resume")
        return
    
    resume_text = response.json()["text"]
    print(f"✅ Uploaded marketing resume ({len(resume_text)} characters)")
    
    # Tech job description
    tech_job = """
    Product Marketing Manager - SaaS Technology
    
    We're seeking a Product Marketing Manager to drive go-to-market strategies for our B2B SaaS platform.
    
    Required Skills:
    • 5+ years in product marketing or digital marketing
    • Experience with SaaS/B2B marketing
    • Strong analytical skills and data-driven decision making
    • Knowledge of marketing automation tools (HubSpot, Marketo)
    • Experience with A/B testing and conversion optimization
    • Understanding of customer lifecycle and retention strategies
    • Excellent communication and presentation skills
    • Experience with product launches and positioning
    
    Responsibilities:
    • Develop product positioning and messaging strategies
    • Create sales enablement materials and training
    • Analyze customer data and market trends
    • Collaborate with product and engineering teams
    • Manage product launch campaigns
    • Optimize conversion funnels and customer onboarding
    """
    
    # Test tailoring
    data = {
        'resume_text': resume_text,
        'job_description': tech_job
    }
    
    response = session.post(f"{BACKEND_URL}/tailor-resume", data=data, timeout=60)
    
    if response.status_code == 200:
        result = response.json()
        ats_score = result.get("ats_score", 0)
        suggestions = result.get("suggestions", [])
        keyword_matches = result.get("keyword_matches", [])
        missing_keywords = result.get("missing_keywords", [])
        tailored_resume = result.get("tailored_resume", "")
        
        print(f"✅ AI Tailoring Successful")
        print(f"   ATS Score: {ats_score}/100")
        print(f"   Suggestions: {len(suggestions)}")
        print(f"   Keyword Matches: {keyword_matches}")
        print(f"   Missing Keywords: {missing_keywords}")
        
        # Check if tailored resume contains tech-relevant terms
        tech_terms = ["saas", "b2b", "product marketing", "automation", "analytics", "conversion"]
        found_terms = [term for term in tech_terms if term.lower() in tailored_resume.lower()]
        print(f"   Tech Terms Added: {found_terms}")
        
        # Validate ATS score is reasonable for a career transition
        if 60 <= ats_score <= 90:
            print("✅ ATS Score is reasonable for career transition")
        else:
            print(f"⚠️  ATS Score ({ats_score}) may be too high/low for career transition")
            
        # Check if suggestions are relevant
        if suggestions and len(suggestions) > 0:
            print("✅ AI provided improvement suggestions")
            for i, suggestion in enumerate(suggestions[:3], 1):
                print(f"   {i}. {suggestion}")
        else:
            print("⚠️  No suggestions provided")
            
    else:
        print(f"❌ AI Tailoring failed: HTTP {response.status_code}")
        print(response.text)
    
    print("\n" + "=" * 60)
    
    # Test 2: Stress test with minimal resume
    print("🧪 Testing AI with Minimal Resume Content")
    print("=" * 60)
    
    minimal_resume = "John Doe\nSoftware Developer\nPython, JavaScript\n2 years experience"
    
    data = {
        'resume_text': minimal_resume,
        'job_description': "Senior Python Developer position requiring 5+ years experience with Django, FastAPI, and cloud technologies."
    }
    
    response = session.post(f"{BACKEND_URL}/tailor-resume", data=data, timeout=60)
    
    if response.status_code == 200:
        result = response.json()
        ats_score = result.get("ats_score", 0)
        tailored_resume = result.get("tailored_resume", "")
        
        print(f"✅ Handled minimal resume")
        print(f"   Original length: {len(minimal_resume)} chars")
        print(f"   Tailored length: {len(tailored_resume)} chars")
        print(f"   ATS Score: {ats_score}/100")
        
        # Check if AI expanded the content appropriately
        if len(tailored_resume) > len(minimal_resume) * 2:
            print("✅ AI appropriately expanded minimal content")
        else:
            print("⚠️  AI may not have expanded content sufficiently")
            
    else:
        print(f"❌ Failed to handle minimal resume: HTTP {response.status_code}")
    
    print("\n" + "=" * 60)
    print("🎯 AI Functionality Testing Complete")

if __name__ == "__main__":
    test_ai_tailoring_scenarios()